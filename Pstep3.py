Phase 2: LLM-Based Whitepaper Analysis and Key Highlights Extraction
In this next phase, we will implement a system where an LLM analyzes the whitepaper, extracts key highlights, and provides a summary of what the document contains. This will help users understand the core content and structure of the whitepaper without reading the entire document.

Implementation Approach
Goals
Extract Key Highlights:

Identify important sections, key findings, and critical insights from the whitepaper.

Summarize methodologies, results, and conclusions.

Document Structure Analysis:

Provide an overview of what the document contains (e.g., sections, tables, figures).

Automated Insights:

Derive actionable insights, such as validation methodologies or assumptions.

Testing Plan:

Ensure accuracy, completeness, and relevance of extracted highlights.

System Architecture
Step 1: Preprocessing
Split the whitepaper into manageable chunks using chunking logic (similar to Phase 1).

Extract metadata (e.g., titles, headers) to identify document structure.

Step 2: LLM Analysis
Use an LLM (e.g., Claude or GPT) to analyze each chunk.

Generate structured outputs for:

Key highlights

Section summaries

Document structure overview

Step 3: Consolidation
Combine chunk-level insights into a single comprehensive summary.

Organize results into structured categories (e.g., Introduction, Methodology, Results).

Step 4: Output Generation
Generate a final report with:

Key highlights

Document structure overview

Actionable insights

Implementation Code
Here’s how we can implement this system:


##################################################
import boto3
import json
import logging
import time
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from docx import Document

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def smart_chunk_text(text, chunk_size=8000, overlap=500):
    """
    Split text into chunks more intelligently, respecting paragraph boundaries.
    
    Args:
        text (str): Full text to chunk.
        chunk_size (int): Target chunk size in characters.
        overlap (int): Overlap between chunks.
        
    Returns:
        list: List of text chunks.
    """
    # Split text by paragraphs (or double newline)
    paragraphs = text.split("\n\n")
    
    chunks = []
    current_chunk = []
    current_size = 0
    
    for paragraph in paragraphs:
        paragraph = paragraph.strip()
        if not paragraph:  # Skip empty paragraphs
            continue
            
        paragraph_size = len(paragraph)
        
        # If adding this paragraph would exceed chunk size and we already have content,
        # finish current chunk and start a new one
        if current_size + paragraph_size > chunk_size and current_chunk:
            chunks.append("\n\n".join(current_chunk))
            
            # Include some overlap from previous chunk
            overlap_size = 0
            overlap_paragraphs = []
            
            # Start from the end and work backwards to add paragraphs for overlap
            for p in reversed(current_chunk):
                if overlap_size + len(p) <= overlap:
                    overlap_paragraphs.insert(0, p)
                    overlap_size += len(p)
                else:
                    break
                    
            current_chunk = overlap_paragraphs
            current_size = overlap_size
        
        # Add the paragraph to the current chunk
        current_chunk.append(paragraph)
        current_size += paragraph_size
    
    # Add the last chunk if it's not empty
    if current_chunk:
        chunks.append("\n\n".join(current_chunk))
    
    return chunks

def validate_json_structure(json_obj):
    """
    Validates that the JSON structure matches the expected format for whitepaper analysis.
    
    Args:
        json_obj: Parsed JSON object to validate
        
    Returns:
        tuple: (is_valid, error_message)
    """
    expected_keys = ["highlights", "summary", "methodologies", "assumptions", "figures_tables"]
    
    # Check if all expected keys exist
    for key in expected_keys:
        if key not in json_obj:
            return False, f"Missing key: {key}"
        
        # Check if the value is a list
        if not isinstance(json_obj[key], list):
            return False, f"Value for '{key}' is not a list"
    
    return True, "Valid JSON structure"

def analyze_chunk_for_highlights(chunk, retry_count=0, max_retries=2):
    """
    Analyzes a single chunk of text using Claude via AWS Bedrock to extract key highlights.
    Implements retry logic and JSON validation.
    
    Args:
        chunk (str): A chunk of text to analyze.
        retry_count (int): Current retry attempt.
        max_retries (int): Maximum number of retry attempts.

    Returns:
        dict: Extracted highlights and structured analysis.
    """
    bedrock_client = boto3.client('bedrock-runtime', region_name='us-east-1')
    
    prompt = f"""
    You are an expert financial model analyst tasked with extracting structured information from a whitepaper section.
    Your job is to analyze the content and output a strictly formatted JSON object.

    ===== TEXT TO ANALYZE =====
    {chunk}
    ===========================

    INSTRUCTIONS:
    1. Carefully read and understand the content.
    2. Extract information into these EXACT categories:
       a. HIGHLIGHTS: Key findings or conclusions (3-7 items)
       b. SUMMARY: Main points in bullet format (5-10 items)
       c. METHODOLOGIES: Technical approaches mentioned (algorithms, techniques)
       d. ASSUMPTIONS: Stated or implied assumptions about data or models
       e. FIGURES_TABLES: References to figures or tables with descriptions

    CRITICAL FORMATTING REQUIREMENTS:
    - Your output MUST be valid JSON with no extra text before or after
    - Each value must be a simple string (no nested objects or dictionaries)
    - Use ONLY these exact keys: "highlights", "summary", "methodologies", "assumptions", "figures_tables"
    - Each key must map to an ARRAY OF STRINGS, even if empty
    - Do not include any explanation text, metadata, or notes outside the JSON structure
    - Never use special characters that could break JSON parsing
    - Ensure all quotes, brackets, and commas are properly placed
    - Escape any quotes within strings using backslash: \\"
    - Never use objects like {{text: "content", importance: "high"}} - use only plain strings

    SAMPLE EXPECTED OUTPUT 1:
    ```
    {{
      "highlights": [
        "The model achieves 95.2% accuracy on the validation dataset.",
        "Linear regression outperformed random forest for this specific application."
      ],
      "summary": [
        "The study compared multiple regression models on financial data.",
        "Linear regression was selected as the optimal approach."
      ],
      "methodologies": [
        "Linear Regression with ridge regularization",
        "Random Forest with 100 estimators"
      ],
      "assumptions": [
        "The data follows a normal distribution.",
        "Features are independent of each other."
      ],
      "figures_tables": [
        "Figure 3: Comparison of model accuracy across different approaches",
        "Table 2: Feature importance rankings"
      ]
    }}
    ```

    SAMPLE EXPECTED OUTPUT 2:
    ```
    {{
      "highlights": [
        "The neural network architecture consists of 3 hidden layers with 64, 32, and 16 neurons respectively."
      ],
      "summary": [
        "The document describes a deep learning approach for financial forecasting.",
        "ReLU activation functions were used throughout the network.",
        "Adam optimizer was selected with a learning rate of 0.001."
      ],
      "methodologies": [
        "Deep neural network with 3 hidden layers",
        "Adam optimization algorithm",
        "Early stopping to prevent overfitting"
      ],
      "assumptions": [
        "The financial time series data is stationary after preprocessing."
      ],
      "figures_tables": []
    }}
    ```

    If a section has no relevant information, use an empty array:
    "methodologies": []

    DO NOT include any text outside the JSON structure. Your response should be ONLY the JSON object.
    """
    
    try:
        response = bedrock_client.invoke_model(
            modelId="anthropic.claude-3-haiku-20240307-v1:0",
            body=json.dumps({
                "messages": [{"role": "user", "content": prompt}],
                "max_tokens": 3000,
                "temperature": 0.2  # Lower temperature for more consistent formatting
            })
        )
        
        response_body = json.loads(response['body'].read().decode('utf-8'))
        content = response_body["content"]
        
        # Extract JSON if it's wrapped in code blocks
        if "```
            json_start = content.find("```json") + 7
            json_end = content.rfind("```
            json_string = content[json_start:json_end].strip()
        elif "```" in content:
            # Just regular code block without language
            json_start = content.find("```
            json_end = content.rfind("```")
            json_string = content[json_start:json_end].strip()
        else:
            json_string = content.strip()
        
        # Parse JSON
        try:
            result = json.loads(json_string)
            
            # Validate structure
            valid, message = validate_json_structure(result)
            if not valid:
                if retry_count < max_retries:
                    logging.warning(f"Invalid JSON structure: {message}. Retrying ({retry_count+1}/{max_retries})...")
                    return analyze_chunk_for_highlights(chunk, retry_count + 1, max_retries)
                else:
                    logging.error(f"Max retries reached. Last error: {message}")
                    return {
                        "highlights": [],
                        "summary": [f"Error processing chunk: {message}"],
                        "methodologies": [],
                        "assumptions": [],
                        "figures_tables": []
                    }
            
            return result
            
        except json.JSONDecodeError as e:
            if retry_count < max_retries:
                logging.warning(f"JSON parsing error: {str(e)}. Retrying ({retry_count+1}/{max_retries})...")
                return analyze_chunk_for_highlights(chunk, retry_count + 1, max_retries)
            else:
                logging.error(f"Failed to parse JSON after {max_retries} attempts. Error: {str(e)}")
                return {
                    "highlights": [],
                    "summary": [f"Error processing chunk: Invalid JSON format"],
                    "methodologies": [],
                    "assumptions": [],
                    "figures_tables": []
                }
    
    except Exception as e:
        logging.error(f"Error processing chunk: {str(e)}")
        return {
            "highlights": [],
            "summary": [f"Error processing chunk: {str(e)}"],
            "methodologies": [],
            "assumptions": [],
            "figures_tables": []
        }

def process_chunks_in_parallel(chunks, max_workers=5, max_retries=2):
    """
    Process multiple chunks in parallel with retry logic.
    
    Args:
        chunks (list): List of text chunks to process.
        max_workers (int): Maximum number of parallel workers.
        max_retries (int): Maximum retry attempts per chunk.
        
    Returns:
        list: Results from all chunks.
    """
    results = []
    processed_count = 0
    failure_count = 0
    
    logging.info(f"Processing {len(chunks)} chunks with {max_workers} parallel workers...")
    start_time = time.time()
    
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        # Submit all tasks
        future_to_chunk = {executor.submit(analyze_chunk_for_highlights, chunk, 0, max_retries): i 
                         for i, chunk in enumerate(chunks)}
        
        # Process results as they complete
        for future in as_completed(future_to_chunk):
            chunk_idx = future_to_chunk[future]
            processed_count += 1
            
            try:
                result = future.result()
                results.append(result)
                
                # Check if this was a failure
                is_error = all(len(result.get(key, [])) == 0 for key in ["highlights", "methodologies", "assumptions", "figures_tables"])
                if is_error and len(result.get("summary", [])) == 1 and "Error" in result.get("summary", [""])[0]:
                    failure_count += 1
                    
                # Log progress
                if processed_count % 5 == 0 or processed_count == len(chunks):
                    logging.info(f"Processed {processed_count}/{len(chunks)} chunks ({failure_count} failures)")
                
            except Exception as e:
                failure_count += 1
                logging.error(f"Chunk {chunk_idx} processing failed with error: {str(e)}")
                # Add empty result with error information
                results.append({
                    "highlights": [],
                    "summary": [f"Failed to process chunk {chunk_idx}: {str(e)}"],
                    "methodologies": [],
                    "assumptions": [],
                    "figures_tables": []
                })
    
    end_time = time.time()
    logging.info(f"Finished processing all chunks in {end_time - start_time:.2f} seconds")
    logging.info(f"Success rate: {((len(chunks) - failure_count) / len(chunks)) * 100:.1f}% ({failure_count} failures out of {len(chunks)} chunks)")
    
    return results

def deduplicate_items(items):
    """
    Remove duplicate items while preserving order.
    
    Args:
        items (list): List of strings to deduplicate.
        
    Returns:
        list: Deduplicated list with order preserved.
    """
    seen = set()
    result = []
    
    for item in items:
        # Normalize the item text for comparison (lowercase, strip whitespace)
        normalized = item.lower().strip()
        if normalized not in seen:
            seen.add(normalized)
            result.append(item)  # Add original item, not normalized version
            
    return result

def consolidate_highlights(results):
    """
    Consolidates highlights from all chunks into a single structured output with deduplication.
    
    Args:
        results (list): List of results from individual chunks.

    Returns:
        dict: Consolidated highlights and structured analysis.
    """
    consolidated = {
        "highlights": [],
        "summary": [],
        "methodologies": [],
        "assumptions": [],
        "figures_tables": []
    }
    
    # First pass - collect all items
    for result in results:
        if result.get("error"):
            logging.warning(f"Skipping failed result: {result['error']}")
            continue
        
        for key in consolidated.keys():
            consolidated[key].extend(result.get(key, []))
    
    # Second pass - deduplicate each category
    for key in consolidated.keys():
        consolidated[key] = deduplicate_items(consolidated[key])
        logging.info(f"After deduplication: {len(consolidated[key])} unique {key}")
    
    return consolidated

def export_to_word_highlights(consolidated_results, output_file_path):
    """
    Exports consolidated highlights to a Word document.
    
    Args:
        consolidated_results (dict): Consolidated structured data from all chunks.
        output_file_path (str): Path to save the Word document.

    Returns:
        None
    """
    doc = Document()
    
    # Add title
    doc.add_heading('Whitepaper Highlights Report', level=1)
    
    # Add highlights section
    doc.add_heading('Key Highlights', level=2)
    for highlight in consolidated_results.get("highlights", []):
        doc.add_paragraph(f"- {highlight}")
    
    # Add summary section
    doc.add_heading('Summary', level=2)
    for bullet_point in consolidated_results.get("summary", []):
        doc.add_paragraph(f"- {bullet_point}")
    
    # Add methodologies section
    doc.add_heading('Methodologies', level=2)
    for methodology in consolidated_results.get("methodologies", []):
        doc.add_paragraph(f"- {methodology}")
    
    # Add assumptions section
    doc.add_heading('Assumptions', level=2)
    for assumption in consolidated_results.get("assumptions", []):
        doc.add_paragraph(f"- {assumption}")
    
    # Add figures/tables section
    doc.add_heading('Figures/Tables', level=2)
    for reference in consolidated_results.get("figures_tables", []):
        doc.add_paragraph(f"- {reference}")
    
    # Save the document
    doc.save(output_file_path)

def main():
    """
    Main workflow for analyzing a whitepaper and extracting key highlights.
    
    Returns:
        None
    """
    # Load extracted text from file
    extracted_text_path = "./extracted_content/extracted_text.txt"
    
    try:
        with open(extracted_text_path, "r", encoding="utf-8") as f:
            extracted_text = f.read()
        
        logging.info(f"Loaded extracted text from {extracted_text_path}")
        
        # Step 1: Chunk the text using smart chunking that respects paragraph boundaries
        logging.info("Smart chunking extracted text...")
        chunks = smart_chunk_text(extracted_text)
        logging.info(f"Created {len(chunks)} chunks")
        
        # Step 2: Analyze chunks in parallel with retry logic
        logging.info("Analyzing chunks in parallel...")
        results = process_chunks_in_parallel(chunks, max_workers=5, max_retries=2)
        
        # Step 3: Consolidate results with deduplication
        logging.info("Consolidating results with deduplication...")
        consolidated_results = consolidate_highlights(results)
        
        # Step 4: Export results to Word document
        output_file_path = "./whitepaper_highlights.docx"
        logging.info(f"Exporting highlights to Word document at {output_file_path}...")
        export_to_word_highlights(consolidated_results, output_file_path)
        
        print(f"Highlights exported successfully to {output_file_path}")
    
    except Exception as e:
        logging.error(f"Error during whitepaper analysis workflow: {str(e)}")

if __name__ == "__main__":
   main()
