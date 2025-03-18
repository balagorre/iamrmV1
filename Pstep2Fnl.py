pip install python-docx

from docx import Document

def export_to_word(final_results, output_file_path):
    """
    Exports consolidated results to a Word document.
    
    Args:
        final_results (dict): Consolidated structured data from all chunks.
        output_file_path (str): Path to save the Word document.
    
    Returns:
        None
    """
    # Create a new Word document
    doc = Document()
    
    # Add title
    doc.add_heading('Whitepaper Analysis Report', level=1)
    
    # Add metadata section
    doc.add_heading('Metadata', level=2)
    metadata = final_results.get("metadata", {})
    doc.add_paragraph(f"Total Chunks: {metadata.get('total_chunks', 0)}")
    doc.add_paragraph(f"Processed Chunks: {metadata.get('processed_chunks', 0)}")
    doc.add_paragraph(f"Failed Chunks: {metadata.get('failed_chunks', 0)}")
    if metadata.get("failed_chunk_ids"):
        doc.add_paragraph(f"Failed Chunk IDs: {', '.join(metadata['failed_chunk_ids'])}")
    
    # Add summary section
    doc.add_heading('Summary', level=2)
    summaries = final_results.get("summary", [])
    for i, summary in enumerate(summaries):
        doc.add_paragraph(f"{i + 1}. {summary}")
    
    # Add inputs section
    doc.add_heading('Inputs', level=2)
    inputs = final_results.get("inputs", [])
    for input_item in inputs:
        name = input_item.get("name", "Unknown")
        description = input_item.get("description", "No description available")
        format_type = input_item.get("format", "Unknown format")
        doc.add_paragraph(f"- {name}: {description} ({format_type})")
    
    # Add outputs section
    doc.add_heading('Outputs', level=2)
    outputs = final_results.get("outputs", [])
    for output_item in outputs:
        name = output_item.get("name", "Unknown")
        description = output_item.get("description", "No description available")
        format_type = output_item.get("format", "Unknown format")
        doc.add_paragraph(f"- {name}: {description} ({format_type})")
    
    # Add calculations section
    doc.add_heading('Calculations', level=2)
    calculations = final_results.get("calculations", [])
    for calculation in calculations:
        doc.add_paragraph(f"- {calculation}")
    
    # Add model performance section
    doc.add_heading('Model Performance', level=2)
    model_performance = final_results.get("model_performance", [])
    for performance_metric in model_performance:
        doc.add_paragraph(f"- {performance_metric}")
    
    # Add solution specification section
    doc.add_heading('Solution Specification', level=2)
    solution_specification = final_results.get("solution_specification", [])
    for spec_item in solution_specification:
        doc.add_paragraph(f"- {spec_item}")
    
    # Add testing summary section
    doc.add_heading('Testing Summary', level=2)
    testing_summary = final_results.get("testing_summary", [])
    for test_case in testing_summary:
        doc.add_paragraph(f"- {test_case}")
    
    # Add reconciliation section
    doc.add_heading('Reconciliation', level=2)
    reconciliation = final_results.get("reconciliation", [])
    for reconciliation_process in reconciliation:
        doc.add_paragraph(f"- {reconciliation_process}")
    
    # Save the document
    doc.save(output_file_path)
    print(f"Document saved to: {output_file_path}")

# Example usage
final_results = {
  "summary": [
      "This chunk describes the model's architecture.",
      "This chunk describes additional details about data preprocessing."
  ],
  "inputs": [
      {"name": "input_1", "description": "Historical sales data used to predict future trends.", "format": "numerical"},
      {"name": "input_2", "description": "Customer segmentation based on demographic data.", "format": "categorical"}
  ],
  "outputs": [
      {"name": "output_1", "description": "Predicted sales figures for the next quarter.", "format": "numerical"},
      {"name": "output_2", "description": "Risk scores for customers based on historical behavior.", "format": "numerical"}
  ],
  "calculations": [
      "Linear regression formula applied to sales data.",
      "Clustering algorithm used for customer segmentation."
  ],
  "model_performance": [
      "Accuracy: The model achieves 95% accuracy on historical sales predictions.",
      "Precision: Precision is measured at 90% for high-risk customer identification."
  ],
  "solution_specification": [
      "Architecture: Microservices-based architecture with components for data ingestion and prediction generation.",
      "Components: Includes modules for feature engineering and real-time prediction serving."
  ],
  "testing_summary": [
      "Test Case_1: Validates input data preprocessing logic.",
      "Test Case_2: Ensures predictions align with expected trends under stable conditions."
  ],
  "reconciliation": [
      "Matches predicted sales figures against actual sales data from previous quarters.",
      "Validates customer risk scores against historical behavior patterns."
  ],
  "metadata": {
      "total_chunks": 10,
      "processed_chunks": 9,
      "failed_chunks": 1,
      "failed_chunk_ids": ["chunk_3"]
  }
}

output_file_path = "./final_whitepaper_analysis.docx"
export_to_word(final_results, output_file_path)
